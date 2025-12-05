#!/usr/bin/env python3
"""
Script para gerar relatório DOCX completo do projeto Uni-WhatsChat
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1):
    """Adiciona cabeçalho com estilo personalizado"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_paragraph(doc, text):
    """Adiciona parágrafo com estilo de código"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return p

def create_report():
    """Cria o relatório completo"""
    doc = Document()
    
    # Configuração da página
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading('Uni-WhatsChat: Sistema de Chat Seguro com Autenticação Mútua e Criptografia', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Trabalho de Segurança Computacional')
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()  # Espaço
    
    # ========== 1. INTRODUÇÃO ==========
    add_heading_with_style(doc, '1. Introdução', 1)
    
    p1 = doc.add_paragraph()
    p1.add_run(
        'O Uni-WhatsChat é um sistema de comunicação em tempo real desenvolvido com foco em segurança '
        'computacional. O sistema implementa múltiplas camadas de proteção para garantir autenticidade, '
        'integridade e confidencialidade das comunicações entre usuários.'
    )
    
    p2 = doc.add_paragraph()
    p2.add_run(
        'Este relatório apresenta uma análise detalhada da arquitetura, implementação e mecanismos de '
        'segurança do sistema, incluindo autenticação mútua via TLS, troca de chaves Diffie-Hellman e '
        'verificação de integridade através de HMAC.'
    )
    
    # ========== 2. OBJETIVOS ==========
    add_heading_with_style(doc, '2. Objetivos', 1)
    
    doc.add_paragraph('O projeto tem como objetivos principais:', style='List Bullet')
    doc.add_paragraph('Implementar autenticação mútua entre servidor e clientes utilizando certificados X.509', style='List Bullet 2')
    doc.add_paragraph('Estabelecer comunicação segura com troca de chaves Diffie-Hellman para Perfect Forward Secrecy', style='List Bullet 2')
    doc.add_paragraph('Garantir integridade das mensagens através de HMAC-SHA256', style='List Bullet 2')
    doc.add_paragraph('Desenvolver um sistema funcional de chat em tempo real com múltiplos usuários', style='List Bullet 2')
    doc.add_paragraph('Demonstrar a aplicação prática de conceitos de segurança computacional', style='List Bullet 2')
    
    # ========== 3. ARQUITETURA DO SISTEMA ==========
    add_heading_with_style(doc, '3. Arquitetura do Sistema', 1)
    
    add_heading_with_style(doc, '3.1. Componentes Principais', 2)
    
    p3 = doc.add_paragraph()
    p3.add_run('O sistema é composto pelos seguintes módulos principais:').bold = True
    
    doc.add_paragraph('Servidor (server.py): Gerencia conexões, autenticação e roteamento de mensagens', style='List Bullet')
    doc.add_paragraph('Cliente (client.py): Interface de usuário e comunicação com o servidor', style='List Bullet')
    doc.add_paragraph('Criptografia (crypto.py): Implementa Diffie-Hellman, HMAC e derivação de chaves', style='List Bullet')
    doc.add_paragraph('Rede (network.py): Configuração de contextos SSL/TLS para mTLS', style='List Bullet')
    doc.add_paragraph('Protocolo (protocol.py): Serialização e deserialização de pacotes de comunicação', style='List Bullet')
    
    add_heading_with_style(doc, '3.2. Fluxo de Comunicação', 2)
    
    p4 = doc.add_paragraph()
    p4.add_run('O fluxo de comunicação segue as seguintes etapas:').bold = True
    
    doc.add_paragraph('1. Conexão TLS: Cliente estabelece conexão TCP e inicia handshake TLS com autenticação mútua', style='List Number')
    doc.add_paragraph('2. Verificação de Certificados: Servidor e cliente verificam mutuamente seus certificados X.509', style='List Number')
    doc.add_paragraph('3. Troca de Chaves Diffie-Hellman: Cliente e servidor trocam chaves públicas e calculam segredo compartilhado', style='List Number')
    doc.add_paragraph('4. Derivação de Chave: Segredo compartilhado é derivado usando HKDF-SHA256', style='List Number')
    doc.add_paragraph('5. Envio de Mensagens: Mensagens são assinadas com HMAC e enviadas via conexão TLS', style='List Number')
    doc.add_paragraph('6. Verificação de Integridade: Servidor verifica HMAC antes de reenviar para outros clientes', style='List Number')
    
    # ========== 4. MECANISMOS DE SEGURANÇA ==========
    add_heading_with_style(doc, '4. Mecanismos de Segurança', 1)
    
    add_heading_with_style(doc, '4.1. Autenticação Mútua (mTLS)', 2)
    
    p5 = doc.add_paragraph()
    p5.add_run(
        'O sistema implementa autenticação mútua utilizando TLS (Transport Layer Security) com certificados X.509. '
        'Diferentemente do TLS tradicional, onde apenas o servidor se autentica, o mTLS exige que tanto o servidor '
        'quanto os clientes apresentem certificados válidos.'
    )
    
    p6 = doc.add_paragraph()
    p6.add_run('Características da implementação:').bold = True
    
    doc.add_paragraph('Autoridade Certificadora (CA) própria: Sistema de PKI com CA central que assina todos os certificados', style='List Bullet')
    doc.add_paragraph('Verificação obrigatória: Servidor rejeita conexões sem certificado válido (CERT_REQUIRED)', style='List Bullet')
    doc.add_paragraph('Identidade baseada em certificado: Username extraído do campo CommonName (CN) do certificado', style='List Bullet')
    doc.add_paragraph('Certificados de 2048 bits: Chaves RSA de tamanho adequado para segurança moderna', style='List Bullet')
    
    p7 = doc.add_paragraph()
    p7.add_run(
        'A implementação utiliza a biblioteca SSL do Python, configurando contextos específicos para servidor e cliente. '
        'O servidor carrega seu certificado e chave privada, além do certificado da CA para verificar clientes. '
        'Os clientes carregam seus próprios certificados e o certificado da CA para verificar o servidor.'
    )
    
    add_heading_with_style(doc, '4.2. Troca de Chaves Diffie-Hellman', 2)
    
    p8 = doc.add_paragraph()
    p8.add_run(
        'O protocolo Diffie-Hellman (DH) é utilizado para estabelecer um segredo compartilhado entre cliente e servidor '
        'sem transmitir a chave diretamente pela rede. Isso proporciona Perfect Forward Secrecy (PFS), garantindo que '
        'mesmo se as chaves de longo prazo forem comprometidas, as comunicações passadas permanecem seguras.'
    )
    
    p9 = doc.add_paragraph()
    p9.add_run('Parâmetros e configuração:').bold = True
    
    doc.add_paragraph('Tamanho da chave: 2048 bits, adequado para segurança atual', style='List Bullet')
    doc.add_paragraph('Gerador: 2 (valor padrão seguro)', style='List Bullet')
    doc.add_paragraph('Parâmetros gerados dinamicamente: Cada servidor gera seus próprios parâmetros p e g', style='List Bullet')
    doc.add_paragraph('Serialização PEM: Chaves públicas serializadas em formato PEM para transmissão', style='List Bullet')
    
    p10 = doc.add_paragraph()
    p10.add_run(
        'Após o cálculo do segredo compartilhado, o sistema utiliza HKDF (HMAC-based Key Derivation Function) com SHA-256 '
        'para derivar uma chave de 32 bytes (256 bits) adequada para uso em HMAC. O HKDF garante que a chave derivada '
        'tenha propriedades criptográficas uniformes, mesmo que o segredo compartilhado tenha alguma estrutura matemática.'
    )
    
    add_heading_with_style(doc, '4.3. Verificação de Integridade (HMAC)', 2)
    
    p11 = doc.add_paragraph()
    p11.add_run(
        'HMAC (Hash-based Message Authentication Code) é utilizado para garantir a integridade e autenticidade das mensagens. '
        'Cada mensagem é assinada com HMAC-SHA256 usando a chave derivada do segredo Diffie-Hellman.'
    )
    
    p12 = doc.add_paragraph()
    p12.add_run('Características da implementação:').bold = True
    
    doc.add_paragraph('Algoritmo: HMAC-SHA256, amplamente reconhecido e seguro', style='List Bullet')
    doc.add_paragraph('Tamanho da assinatura: 64 caracteres hexadecimais (256 bits)', style='List Bullet')
    doc.add_paragraph('Comparação segura: Utiliza hmac.compare_digest() para evitar timing attacks', style='List Bullet')
    doc.add_paragraph('Verificação obrigatória: Mensagens sem HMAC válido são rejeitadas', style='List Bullet')
    
    p13 = doc.add_paragraph()
    p13.add_run(
        'O processo funciona da seguinte forma: o remetente calcula o HMAC da mensagem usando a chave compartilhada e '
        'anexa o resultado ao pacote. O receptor recalcula o HMAC e compara com o recebido. Se houver qualquer diferença, '
        'a mensagem é rejeitada, indicando possível modificação ou ataque.'
    )
    
    # ========== 5. IMPLEMENTAÇÃO TÉCNICA ==========
    add_heading_with_style(doc, '5. Implementação Técnica', 1)
    
    add_heading_with_style(doc, '5.1. Protocolo de Comunicação', 2)
    
    p14 = doc.add_paragraph()
    p14.add_run(
        'O protocolo de comunicação utiliza JSON para serialização de dados. Cada pacote é uma linha JSON terminada por '
        'newline (\\n), facilitando o parsing e permitindo streaming de mensagens.'
    )
    
    p15 = doc.add_paragraph()
    p15.add_run('Tipos de pacote:').bold = True
    
    doc.add_paragraph('KEY_EXCHANGE: Contém chave pública Diffie-Hellman codificada em Base64', style='List Bullet')
    doc.add_paragraph('MSG: Contém mensagem de chat com remetente, conteúdo e assinatura HMAC', style='List Bullet')
    
    p16 = doc.add_paragraph()
    p16.add_run('Estrutura do pacote MSG:').bold = True
    add_code_paragraph(doc, '{\n  "type": "MSG",\n  "sender": "alice",\n  "content": "Olá!",\n  "hmac": "a1b2c3..."\n}')
    
    add_heading_with_style(doc, '5.2. Gerenciamento de Conexões', 2)
    
    p17 = doc.add_paragraph()
    p17.add_run(
        'O servidor utiliza threading para gerenciar múltiplas conexões simultâneas. Cada cliente é atendido em uma thread '
        'separada, permitindo comunicação assíncrona. O servidor mantém um dicionário de clientes conectados, armazenando '
        'informações como socket, chave compartilhada e chave privada DH.'
    )
    
    p18 = doc.add_paragraph()
    p18.add_run(
        'O sistema implementa broadcast de mensagens: quando um cliente envia uma mensagem, o servidor verifica a integridade '
        'e reenvia para todos os outros clientes conectados, criando um ambiente de chat em grupo.'
    )
    
    add_heading_with_style(doc, '5.3. Geração de Certificados', 2)
    
    p19 = doc.add_paragraph()
    p19.add_run(
        'O sistema inclui scripts para geração automatizada de certificados. O script setup_certs.py cria a CA e o certificado '
        'do servidor, enquanto create_user.py gera certificados para cada usuário, todos assinados pela mesma CA.'
    )
    
    p20 = doc.add_paragraph()
    p20.add_run('Processo de geração:').bold = True
    
    doc.add_paragraph('1. Geração de chave privada RSA de 2048 bits', style='List Number')
    doc.add_paragraph('2. Criação de Certificate Signing Request (CSR)', style='List Number')
    doc.add_paragraph('3. Assinatura pela CA usando SHA-256', style='List Number')
    doc.add_paragraph('4. Validade de 365 dias para certificados de usuário', style='List Number')
    
    # ========== 6. ANÁLISE DE SEGURANÇA ==========
    add_heading_with_style(doc, '6. Análise de Segurança', 1)
    
    add_heading_with_style(doc, '6.1. Forças do Sistema', 2)
    
    p21 = doc.add_paragraph()
    p21.add_run('O sistema apresenta várias características de segurança robustas:').bold = True
    
    doc.add_paragraph('Autenticação mútua garante que apenas usuários autorizados possam se conectar', style='List Bullet')
    doc.add_paragraph('Perfect Forward Secrecy protege comunicações passadas mesmo com comprometimento futuro de chaves', style='List Bullet')
    doc.add_paragraph('Verificação de integridade detecta qualquer modificação nas mensagens', style='List Bullet')
    doc.add_paragraph('Uso de algoritmos criptográficos modernos e amplamente testados', style='List Bullet')
    doc.add_paragraph('Implementação seguindo boas práticas de segurança (comparação segura de HMAC, derivação adequada de chaves)', style='List Bullet')
    
    add_heading_with_style(doc, '6.2. Limitações e Melhorias Futuras', 2)
    
    p22 = doc.add_paragraph()
    p22.add_run('O sistema atual possui algumas limitações que podem ser melhoradas:').bold = True
    
    doc.add_paragraph('Parâmetros DH: Idealmente, o servidor deveria enviar os parâmetros DH para os clientes, garantindo uso dos mesmos parâmetros', style='List Bullet')
    doc.add_paragraph('Criptografia de conteúdo: Atualmente, apenas a integridade é verificada; adicionar criptografia AES das mensagens aumentaria a confidencialidade', style='List Bullet')
    doc.add_paragraph('Revogação de certificados: Sistema não implementa CRL (Certificate Revocation List) ou OCSP', style='List Bullet')
    doc.add_paragraph('Renegociação de chaves: Não há renovação periódica das chaves DH durante uma sessão longa', style='List Bullet')
    doc.add_paragraph('Proteção contra replay attacks: Sistema não implementa nonces ou timestamps para prevenir replay', style='List Bullet')
    
    add_heading_with_style(doc, '6.3. Ameaças Mitigadas', 2)
    
    p23 = doc.add_paragraph()
    p23.add_run('O sistema mitiga as seguintes ameaças:').bold = True
    
    doc.add_paragraph('Man-in-the-Middle (MITM): Autenticação mútua e verificação de certificados impedem ataques MITM', style='List Bullet')
    doc.add_paragraph('Modificação de mensagens: HMAC detecta qualquer alteração no conteúdo', style='List Bullet')
    doc.add_paragraph('Impersonação: Certificados garantem identidade dos usuários', style='List Bullet')
    doc.add_paragraph('Eavesdropping: Conexão TLS criptografa todo o tráfego', style='List Bullet')
    
    # ========== 7. CONCLUSÃO ==========
    add_heading_with_style(doc, '7. Conclusão', 1)
    
    p24 = doc.add_paragraph()
    p24.add_run(
        'O Uni-WhatsChat demonstra a aplicação prática de conceitos fundamentais de segurança computacional em um sistema '
        'funcional de comunicação. A implementação combina autenticação mútua, troca de chaves segura e verificação de '
        'integridade para criar um ambiente de comunicação robusto.'
    )
    
    p25 = doc.add_paragraph()
    p25.add_run(
        'O projeto serve como uma excelente base educacional para compreender como diferentes mecanismos de segurança podem '
        'ser integrados para proteger comunicações. Embora existam oportunidades de melhoria, o sistema atual atende aos '
        'objetivos de demonstrar autenticidade, integridade e confidencialidade em um ambiente prático.'
    )
    
    p26 = doc.add_paragraph()
    p26.add_run(
        'Futuras melhorias podem incluir criptografia de conteúdo adicional, sistema de revogação de certificados, e interface '
        'gráfica para melhor experiência do usuário. No entanto, a arquitetura atual fornece uma base sólida para essas '
        'expansões.'
    )
    
    # ========== 8. REFERÊNCIAS ==========
    add_heading_with_style(doc, '8. Referências', 1)
    
    doc.add_paragraph('Dierks, T., & Rescorla, E. (2018). The Transport Layer Security (TLS) Protocol Version 1.3. RFC 8446.', style='List Bullet')
    doc.add_paragraph('Krawczyk, H., & Eronen, P. (2010). HMAC-based Extract-and-Expand Key Derivation Function (HKDF). RFC 5869.', style='List Bullet')
    doc.add_paragraph('Krawczyk, H., Bellare, M., & Canetti, R. (1997). HMAC: Keyed-Hashing for Message Authentication. RFC 2104.', style='List Bullet')
    doc.add_paragraph('Kivinen, T., & Kojo, M. (2003). More Modular Exponential (MODP) Diffie-Hellman groups for Internet Key Exchange (IKE). RFC 3526.', style='List Bullet')
    doc.add_paragraph('Python Software Foundation. (2024). ssl — TLS/SSL wrapper for socket objects. Python Documentation.', style='List Bullet')
    doc.add_paragraph('The Python Cryptographic Authority. (2024). cryptography. https://cryptography.io/', style='List Bullet')
    
    # Salva o documento
    output_file = 'Relatorio_Uni_WhatsChat.docx'
    doc.save(output_file)
    print(f"Relatório gerado com sucesso: {output_file}")

if __name__ == "__main__":
    create_report()


